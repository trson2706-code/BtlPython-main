"""Tests cho ReportGenerator — Ch1-6 (full 6-chapter report)."""

import os
import pytest
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from src.core.report_generator import ReportGenerator


@pytest.fixture
def report_path(tmp_path):
    """Tạo report trong thư mục tạm."""
    path = str(tmp_path / 'bao_cao.docx')
    gen = ReportGenerator(output_path=path)
    gen.generate()
    return path


@pytest.fixture
def doc(report_path):
    """Load document đã tạo."""
    return Document(report_path)


# Test 1: File tạo thành công
def test_file_created(report_path):
    assert os.path.exists(report_path)
    assert os.path.getsize(report_path) > 0


# Test 2: Chương 1 heading
def test_chapter_1_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 1' in h for h in headings)


# Test 3: Chương 2 heading
def test_chapter_2_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 2' in h for h in headings)


# Test 4: Bảng so sánh thư viện (≥3 cột, ≥6 dòng)
def test_library_comparison_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Thư viện' in headers and 'Version' in headers:
            assert len(headers) >= 3
            assert len(table.rows) >= 7  # 1 header + 6 data
            return
    pytest.fail("Bảng so sánh thư viện không tìm thấy")


# Test 5: Bảng face_recognition vs DeepFace (≥7 tiêu chí, 3-col discriminator — F2-fix)
def test_face_recognition_vs_deepface_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        # F2-fix: require 3 columns to avoid matching Ch5 2-col test results table
        if 'Tiêu chí' in headers and len(headers) == 3:
            cells_col0 = [table.rows[i].cells[0].text for i in range(1, len(table.rows))]
            assert len(cells_col0) >= 7
            return
    pytest.fail("Bảng so sánh face_recognition vs DeepFace không tìm thấy (cần 3-col with 'Tiêu chí')")


# Test 6: Margins đúng
def test_document_margins(doc):
    section = doc.sections[0]
    tolerance = 1000  # EMU rounding tolerance
    assert abs(section.left_margin - Cm(3)) < tolerance
    assert abs(section.right_margin - Cm(2)) < tolerance
    assert abs(section.top_margin - Cm(2)) < tolerance
    assert abs(section.bottom_margin - Cm(2)) < tolerance


# Test 7: Body text font size = 13pt
def test_body_font_size(doc):
    style = doc.styles['Normal']
    assert style.font.size == Pt(13)


# Test 8: Flowchart section tồn tại
def test_flowchart_section(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('2.4' in h or 'Flowchart' in h or 'flowchart' in h.lower() for h in headings)


# Test 9: Database schema table có 6 rows data
def test_database_schema_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Bảng' in headers and 'Mô tả' in headers:
            data_rows = len(table.rows) - 1
            assert data_rows == 6, f"Expected 6 data rows, got {data_rows}"
            return
    pytest.fail("Bảng database schema không tìm thấy")


# Test 10: Idempotent - overwrite khi chạy lại (F8-fix: content-based check)
def test_overwrite_on_rerun(tmp_path):
    path = str(tmp_path / 'bao_cao.docx')
    gen = ReportGenerator(output_path=path)
    gen.generate()
    size1 = os.path.getsize(path)
    gen.generate()
    size2 = os.path.getsize(path)
    # File should still exist and be valid after second run
    assert os.path.exists(path)
    assert size2 > 0
    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 1' in h for h in headings)
    # Sizes should be equal (idempotent content)
    assert size1 == size2


# Test 11: Trang bìa chứa text
def test_cover_page(doc):
    all_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Đề tài 15' in all_text
    assert 'TRƯỜNG ĐẠI HỌC' in all_text


# Test 12: Page numbering - footer có PAGE field
def test_page_numbering(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_paras = footer.paragraphs
    assert len(footer_paras) > 0
    # Check footer XML contains PAGE field
    footer_xml = footer._element.xml
    assert 'PAGE' in footer_xml


# Test 13: Line spacing = 1.5
def test_line_spacing(doc):
    style = doc.styles['Normal']
    assert style.paragraph_format.line_spacing == 1.5


# Test 14: Heading font = Times New Roman
def test_heading_font(doc):
    h1_style = doc.styles['Heading 1']
    assert h1_style.font.name == 'Times New Roman'
    h2_style = doc.styles['Heading 2']
    assert h2_style.font.name == 'Times New Roman'


# Test 15: Vietnamese text rendering
def test_vietnamese_text(doc):
    all_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'khuôn mặt' in all_text
    assert 'điểm danh' in all_text
    assert 'sinh viên' in all_text


# ── NEW TESTS (F7-F11 remediations) ──


# Test 16: F7-fix — Empty output path raises ValueError
def test_empty_output_path_raises():
    with pytest.raises(ValueError, match="output_path"):
        ReportGenerator(output_path='')


# Test 17: F7-fix — Whitespace-only output path raises ValueError
def test_whitespace_output_path_raises():
    with pytest.raises(ValueError, match="output_path"):
        ReportGenerator(output_path='   ')


# Test 18: F9-fix — Heading sizes (H1=16pt, H2=14pt)
def test_heading_sizes(doc):
    h1_style = doc.styles['Heading 1']
    assert h1_style.font.size == Pt(16), f"H1 size: {h1_style.font.size}"
    h2_style = doc.styles['Heading 2']
    assert h2_style.font.size == Pt(14), f"H2 size: {h2_style.font.size}"


# Test 19: F10-fix — Body text alignment = JUSTIFY
def test_body_text_justify(doc):
    style = doc.styles['Normal']
    assert style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


# Test 20: F11-fix — Table header cells have shading
def test_table_header_shading(doc):
    """Verify at least one table has header cells with background shading."""
    found_shading = False
    for table in doc.tables:
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            tcPr = cell._element.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None and shd.get(qn('w:fill')):
                    found_shading = True
                    break
        if found_shading:
            break
    assert found_shading, "No table header shading found"


# ── NEW TESTS (Party-mode review remediations) ──


# Test 21: F4-fix — Heading color = black (not theme blue)
def test_heading_color_black(doc):
    h1_style = doc.styles['Heading 1']
    assert h1_style.font.color.rgb == RGBColor(0, 0, 0), \
        f"H1 color: {h1_style.font.color.rgb}, expected black"
    # Also verify actual heading runs have black color
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            for run in p.runs:
                assert run.font.color.rgb == RGBColor(0, 0, 0), \
                    f"Heading 1 run color: {run.font.color.rgb}"
            break


# Test 22: F5-fix — East Asian font set on Normal style
def test_east_asian_font(doc):
    style = doc.styles['Normal']
    rFonts = style.element.rPr.rFonts
    ea_font = rFonts.get(qn('w:eastAsia'))
    assert ea_font == 'Times New Roman', f"East Asian font: {ea_font}"


# Test 23: F6-fix — Cover page followed by page break
def test_cover_page_break(doc):
    """Cover page should end with a page break before Chapter 1."""
    found_break = False
    for p in doc.paragraphs:
        if 'Chương 1' in p.text:
            break
        # Check for page break in paragraph runs
        for run in p.runs:
            if run._element.xml and 'w:br' in run._element.xml and 'page' in run._element.xml:
                found_break = True
    # Also check paragraph XML for lastRenderedPageBreak or br type="page"
    if not found_break:
        for p in doc.paragraphs:
            if 'Chương 1' in p.text:
                break
            xml = p._element.xml
            if 'w:br' in xml and 'page' in xml:
                found_break = True
    assert found_break, "No page break found after cover page"


# Test 24: F3-fix — None output path raises ValueError
def test_none_output_path_raises():
    with pytest.raises(ValueError, match="output_path"):
        ReportGenerator(output_path=None)


# Test 25: F10-fix — Subdirectory auto-creation
def test_subdirectory_auto_creation(tmp_path):
    nested_path = str(tmp_path / 'sub' / 'deep' / 'bao_cao.docx')
    gen = ReportGenerator(output_path=nested_path)
    gen.generate()
    assert os.path.exists(nested_path)
    doc = Document(nested_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 1' in h for h in headings)


# Test 26: F12-fix — Invalid save path raises OSError
def test_invalid_save_path_raises(tmp_path):
    """Saving to a path where parent is a file should raise OSError."""
    # Create a file that blocks directory creation
    blocker = tmp_path / 'blocker'
    blocker.write_text('block')
    bad_path = str(blocker / 'sub' / 'bao_cao.docx')
    gen = ReportGenerator(output_path=bad_path)
    with pytest.raises(OSError):
        gen.generate()


# ── TESTS 27-38: Chương 3-4 (7-2-report-technique-model) ──


def _text_after_heading(doc, chapter_keyword):
    """Return concatenated text of paragraphs after the chapter heading.

    Uses 'Heading 1' style boundary detection: stops at next H1 whose text
    does NOT contain chapter_keyword. This prevents cross-chapter leakage.
    """
    found = False
    texts = []
    for p in doc.paragraphs:
        if found and p.style.name == 'Heading 1' and chapter_keyword not in p.text:
            break  # Reached next chapter
        if chapter_keyword in p.text and p.style.name == 'Heading 1':
            found = True
            continue
        if found:
            texts.append(p.text)
    return '\n'.join(texts)


# Test 27: Chương 3 heading
def test_chapter_3_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 3' in h for h in headings)


# Test 28: Chương 4 heading
def test_chapter_4_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 4' in h for h in headings)


# Test 29: Bảng thư viện cài đặt (≥9 rows data + package name validation)
def test_installation_packages_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        # Ch3 table has exactly 3 cols: Thư viện, Version, Vai trò
        # Ch1 table has 5 cols: Thư viện, Version, Vai trò, Ưu điểm, Nhược điểm
        # len==3 discriminates between the two (F4-fix: documented)
        if 'Thư viện' in headers and 'Vai trò' in headers and len(headers) == 3:
            assert len(table.rows) >= 10, f"Expected ≥10 rows, got {len(table.rows)}"
            # Validate specific package names exist
            all_cells = [table.rows[r].cells[0].text for r in range(1, len(table.rows))]
            for pkg in ['opencv-python', 'face-recognition', 'dlib-bin']:
                assert any(pkg in c for c in all_cells), f"Missing package: {pkg}"
            return
    pytest.fail("Bảng thư viện cài đặt không tìm thấy")


# Test 30: Bảng config params — STRICT table check, no text fallback
def test_config_params_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Tham số' in headers and 'Giá trị' in headers:
            return
    pytest.fail("Bảng config params không tìm thấy (cần header 'Tham số' + 'Giá trị')")


# Test 31: Pipeline text chứa HOG — SCOPED to Chapter 4
def test_chapter_4_hog_content(doc):
    ch4_text = _text_after_heading(doc, 'Chương 4')
    assert 'HOG' in ch4_text, "Chapter 4 missing HOG content"


# Test 32: ResNet content — SCOPED to Chapter 4 (F2-fix: assert both independently)
def test_chapter_4_resnet_content(doc):
    ch4_text = _text_after_heading(doc, 'Chương 4')
    assert 'ResNet' in ch4_text, "Chapter 4 missing ResNet keyword"
    assert '128' in ch4_text, "Chapter 4 missing 128d vector reference"


# Test 33: Euclidean content — SCOPED to Chapter 4 (F3-fix: assert Euclidean specifically)
def test_chapter_4_euclidean_content(doc):
    ch4_text = _text_after_heading(doc, 'Chương 4')
    assert 'Euclidean' in ch4_text, "Chapter 4 missing Euclidean keyword"


# Test 34: Bảng thông số kỹ thuật
def test_technical_specs_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Thông số' in headers and 'Giá trị' in headers:
            return
    pytest.fail("Bảng thông số kỹ thuật không tìm thấy")


# Test 35: Backward compat — headings + tables from Ch1-2
def test_backward_compat_full(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 1' in h for h in headings)
    assert any('Chương 2' in h for h in headings)
    # Tables from Ch1-2 must still exist
    found_lib = found_db = False
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Thư viện' in headers and 'Version' in headers and 'Ưu điểm' in headers:
            found_lib = True
        if 'Bảng' in headers and 'Mô tả' in headers:
            found_db = True
    assert found_lib, "Library comparison table from Ch1 missing"
    assert found_db, "Database schema table from Ch2 missing"


# Test 36: Total heading count ≥ 27 (4 H1 + 23 H2)
def test_total_heading_count(doc):
    headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert len(headings) >= 27, f"Expected ≥27 headings, got {len(headings)}"


# Test 37: Bullet lists exist in Chapter 3
def test_chapter_3_bullet_lists(doc):
    found_ch3 = False
    for p in doc.paragraphs:
        if 'Chương 3' in p.text and p.style.name.startswith('Heading'):
            found_ch3 = True
        if found_ch3 and 'Chương 4' in p.text and p.style.name.startswith('Heading'):
            break
        if found_ch3 and p.style.name == 'List Bullet':
            return
    pytest.fail("No bullet lists found in Chapter 3")


# Test 38: Config params table has ≥12 rows (including paths)
def test_config_params_count(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Tham số' in headers and 'Giá trị' in headers:
            assert len(table.rows) >= 13, \
                f"Config table needs ≥13 rows (1 header + 12 data), got {len(table.rows)}"
            return
    pytest.fail("Bảng config params không tìm thấy")


# Test 39: F6-fix — Ch3 section 3.3 Python 3.13 pkg_resources content
def test_chapter_3_python313_error_content(doc):
    ch3_text = _text_after_heading(doc, 'Chương 3')
    assert 'pkg_resources' in ch3_text, "Chapter 3 missing pkg_resources error content"
    assert 'Python 3.13' in ch3_text or '3.13' in ch3_text, \
        "Chapter 3 missing Python 3.13 reference"


# Test 40: F7-fix — Ch4 section 4.3 Face Landmark 68 points content
def test_chapter_4_landmark_content(doc):
    ch4_text = _text_after_heading(doc, 'Chương 4')
    assert '68' in ch4_text, "Chapter 4 missing 68 landmarks reference"
    assert 'landmark' in ch4_text.lower() or 'landmarks' in ch4_text.lower(), \
        "Chapter 4 missing landmark keyword"


# Test 41: F8-fix — Bullet lists exist in Chapter 4
def test_chapter_4_bullet_lists(doc):
    found_ch4 = False
    for p in doc.paragraphs:
        if 'Chương 4' in p.text and p.style.name == 'Heading 1':
            found_ch4 = True
        if found_ch4 and p.style.name == 'Heading 1' and 'Chương 4' not in p.text:
            break  # Reached next chapter
        if found_ch4 and p.style.name == 'List Bullet':
            return
    pytest.fail("No bullet lists found in Chapter 4")


# Test 42: F10-fix — Pipeline table (5 steps) exists in Chapter 4
def test_pipeline_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Bước' in headers and 'Tên' in headers and 'Mô tả' in headers:
            data_rows = len(table.rows) - 1
            assert data_rows == 5, f"Pipeline table needs 5 data rows, got {data_rows}"
            return
    pytest.fail("Bảng pipeline 5 bước không tìm thấy")


# Test 43: Chapter 3 has exactly 5 H2 sub-sections (3.1-3.5)
def test_chapter_3_section_count(doc):
    ch3_sections = []
    in_ch3 = False
    for p in doc.paragraphs:
        if 'Chương 3' in p.text and p.style.name == 'Heading 1':
            in_ch3 = True
            continue
        if in_ch3 and p.style.name == 'Heading 1':
            break
        if in_ch3 and p.style.name == 'Heading 2':
            ch3_sections.append(p.text)
    assert len(ch3_sections) == 5, \
        f"Chapter 3 should have 5 sections (3.1-3.5), got {len(ch3_sections)}: {ch3_sections}"


# Test 44: F4-fix — Chapter 4 has exactly 7 H2 sub-sections (4.1-4.7)
def test_chapter_4_section_count(doc):
    ch4_sections = []
    in_ch4 = False
    for p in doc.paragraphs:
        if 'Chương 4' in p.text and p.style.name == 'Heading 1':
            in_ch4 = True
            continue
        if in_ch4 and p.style.name == 'Heading 1':
            break
        if in_ch4 and p.style.name == 'Heading 2':
            ch4_sections.append(p.text)
    assert len(ch4_sections) == 7, \
        f"Chapter 4 should have 7 sections (4.1-4.7), got {len(ch4_sections)}: {ch4_sections}"


# Test 45: F5-fix — Euclidean formula content exists
def test_euclidean_formula_content(doc):
    ch4_text = _text_after_heading(doc, 'Chương 4')
    assert '√' in ch4_text or 'Σ' in ch4_text, \
        "Chapter 4 missing Euclidean formula (√Σ notation)"


# ── TESTS 46-60: Chương 5-6 (7-3-report-results-conclusion) ──


# Test 46: Chương 5 heading
def test_chapter_5_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 5' in h for h in headings)


# Test 47: Chương 6 heading
def test_chapter_6_heading(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert any('Chương 6' in h for h in headings)


# Test 48: GUI content — SCOPED to Chapter 5
def test_chapter_5_gui_content(doc):
    ch5_text = _text_after_heading(doc, 'Chương 5')
    assert 'giao diện' in ch5_text.lower(), "Chapter 5 missing 'giao diện' keyword"
    # Also verify GUI framework mentioned
    assert 'CustomTkinter' in ch5_text or 'dark' in ch5_text.lower(), \
        "Chapter 5 missing GUI framework reference"


# Test 49: Điểm danh content — SCOPED to Chapter 5
def test_chapter_5_attendance_content(doc):
    ch5_text = _text_after_heading(doc, 'Chương 5')
    assert 'điểm danh' in ch5_text, "Chapter 5 missing điểm danh content"


# Test 50: Excel content — SCOPED to Chapter 5
def test_chapter_5_excel_content(doc):
    ch5_text = _text_after_heading(doc, 'Chương 5')
    assert 'Excel' in ch5_text, "Chapter 5 missing Excel content"


# Test 51: Bảng kết quả thử nghiệm (2-col discriminator + row count — F4-fix)
def test_test_results_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Tiêu chí' in headers and 'Kết quả' in headers and len(headers) == 2:
            # F4-fix: verify row count (8 criteria + 1 header = 9 rows)
            assert len(table.rows) >= 9, \
                f"Test results table needs ≥9 rows (1 header + 8 data), got {len(table.rows)}"
            return
    pytest.fail("Bảng kết quả thử nghiệm không tìm thấy (cần 2-col: 'Tiêu chí' + 'Kết quả')")


# Test 52: Bảng đánh giá hiệu năng (3-col discriminator + row count — F7-fix)
def test_performance_evaluation_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Đánh giá' in headers and len(headers) == 3:
            # F7-fix: verify row count (6 metrics + 1 header = 7 rows)
            assert len(table.rows) >= 7, \
                f"Performance table needs ≥7 rows (1 header + 6 data), got {len(table.rows)}"
            return
    pytest.fail("Bảng đánh giá hiệu năng không tìm thấy (cần 3-col with 'Đánh giá')")


# Test 53: Kết quả đạt được — SCOPED to Chapter 6
def test_chapter_6_achievements(doc):
    ch6_text = _text_after_heading(doc, 'Chương 6')
    assert 'đạt được' in ch6_text.lower(), "Chapter 6 missing achievements content"


# Test 54: Hạn chế — SCOPED to Chapter 6
def test_chapter_6_limitations(doc):
    ch6_text = _text_after_heading(doc, 'Chương 6')
    assert 'hạn chế' in ch6_text.lower(), "Chapter 6 missing limitations content"


# Test 55: Hướng phát triển — SCOPED to Chapter 6
def test_chapter_6_future_work(doc):
    ch6_text = _text_after_heading(doc, 'Chương 6')
    assert 'phát triển' in ch6_text.lower(), "Chapter 6 missing future work content"


# Test 56: Backward compat — Ch1-4 headings + key tables
def test_backward_compat_all_chapters(doc):
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    for ch in ['Chương 1', 'Chương 2', 'Chương 3', 'Chương 4']:
        assert any(ch in h for h in headings), f"{ch} heading missing"
    # Key tables from Ch1-4 must still exist
    found_lib = found_db = found_config = False
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Thư viện' in headers and 'Ưu điểm' in headers:
            found_lib = True
        if 'Bảng' in headers and 'Mô tả' in headers:
            found_db = True
        if 'Tham số' in headers and 'Giá trị' in headers:
            found_config = True
    assert found_lib, "Library comparison table from Ch1 missing"
    assert found_db, "Database schema table from Ch2 missing"
    assert found_config, "Config params table from Ch3 missing"


# Test 57: Total heading count ≥ 35 (6 H1 + 29 H2 minimum)
def test_total_heading_count_6ch(doc):
    headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert len(headings) >= 35, f"Expected ≥35 headings, got {len(headings)}"


# Test 58: Chapter 5 has exactly 5 H2 sub-sections (5.1-5.5)
def test_chapter_5_section_count(doc):
    ch5_sections = []
    in_ch5 = False
    for p in doc.paragraphs:
        if 'Chương 5' in p.text and p.style.name == 'Heading 1':
            in_ch5 = True
            continue
        if in_ch5 and p.style.name == 'Heading 1':
            break
        if in_ch5 and p.style.name == 'Heading 2':
            ch5_sections.append(p.text)
    assert len(ch5_sections) == 5, \
        f"Chapter 5 should have 5 sections (5.1-5.5), got {len(ch5_sections)}: {ch5_sections}"


# Test 59: Chapter 6 has exactly 3 H2 sub-sections (6.1-6.3)
def test_chapter_6_section_count(doc):
    ch6_sections = []
    in_ch6 = False
    for p in doc.paragraphs:
        if 'Chương 6' in p.text and p.style.name == 'Heading 1':
            in_ch6 = True
            continue
        if in_ch6 and p.style.name == 'Heading 1':
            break
        if in_ch6 and p.style.name == 'Heading 2':
            ch6_sections.append(p.text)
    assert len(ch6_sections) == 3, \
        f"Chapter 6 should have 3 sections (6.1-6.3), got {len(ch6_sections)}: {ch6_sections}"


# Test 60: Bullet lists exist in Chapter 6 (limitations/future work)
def test_chapter_6_bullet_lists(doc):
    found_ch6 = False
    for p in doc.paragraphs:
        if 'Chương 6' in p.text and p.style.name == 'Heading 1':
            found_ch6 = True
        if found_ch6 and p.style.name == 'Heading 1' and 'Chương 6' not in p.text:
            break  # Reached next chapter or end
        if found_ch6 and p.style.name == 'List Bullet':
            return
    pytest.fail("No bullet lists found in Chapter 6")


# Test 61: F5-fix — Ch6 summary table (3-col: Mục tiêu, Kết quả, Trạng thái) + row count
def test_chapter_6_summary_table(doc):
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        if 'Mục tiêu' in headers and 'Trạng thái' in headers and len(headers) == 3:
            assert len(table.rows) >= 7, \
                f"Ch6 summary table needs ≥7 rows (1 header + 6 data), got {len(table.rows)}"
            return
    pytest.fail("Bảng tóm tắt kết quả Ch6 không tìm thấy (cần 3-col: 'Mục tiêu' + 'Kết quả' + 'Trạng thái')")


# Test 62: F6-fix — Bullet lists exist in Chapter 5
def test_chapter_5_bullet_lists(doc):
    found_ch5 = False
    for p in doc.paragraphs:
        if 'Chương 5' in p.text and p.style.name == 'Heading 1':
            found_ch5 = True
        if found_ch5 and p.style.name == 'Heading 1' and 'Chương 5' not in p.text:
            break  # Reached next chapter
        if found_ch5 and p.style.name == 'List Bullet':
            return
    pytest.fail("No bullet lists found in Chapter 5")
